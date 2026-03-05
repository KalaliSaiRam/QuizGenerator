from docx import Document
import io

def extract_docx_text(file_bytes):

    doc = Document(io.BytesIO(file_bytes))

    parts = []

    for para in doc.paragraphs:

        if para.text.strip():

            parts.append(para.text.strip())

    for table in doc.tables:

        for row in table.rows:

            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )

            if row_text:

                parts.append(row_text)

    return "\n".join(parts)