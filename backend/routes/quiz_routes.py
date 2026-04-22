from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Header
from ai_engine.engine import process_input, generate_questions, analyze_performance
from pydantic import BaseModel
import fitz
from docx import Document
from PIL import Image
import pytesseract
import tempfile
import os
import pymongo
import datetime
import bcrypt
from dotenv import load_dotenv

load_dotenv()

# Database setup
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017")
try:
    mongo_client = pymongo.MongoClient(MONGO_URI)
    db = mongo_client["quiz_database"]
    history_collection = db["history"]
    users_collection = db["users"]
except Exception as e:
    print(f"Error connecting to MongoDB: {e}")
    history_collection = None
    users_collection = None

import platform

if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# On Linux (Render), tesseract should be in the system path so we don't set tesseract_cmd
router = APIRouter()

# ----------------------------
# Enhanced PDF Extraction (with OCR Fallback)
# ----------------------------
def extract_pdf(file_path):
    doc = fitz.open(file_path)
    pages = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text()

        if not text.strip():
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img, config="--oem 3 --psm 6")
        
        if text.strip():
            pages.append(text)

    doc.close()
    return "\n".join(pages)

# ----------------------------
# DOCX Extraction
# ----------------------------
def extract_docx(file_path):
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)

# ----------------------------
# Enhanced Image OCR Extraction
# ----------------------------
def extract_image(file_path):
    img = Image.open(file_path)
    if img.mode != "RGB":
        img = img.convert("RGB")
    text = pytesseract.image_to_string(img, config="--oem 3 --psm 6")
    return text

# ----------------------------
# Adaptive Feature Helpers
# ----------------------------
def get_user_performance_stats(user_id: str):
    if history_collection is None or not user_id:
        return None
    try:
        rows = list(history_collection.find({"user_id": user_id}, {"topic": 1, "percentage": 1, "_id": 0}))
        if not rows: return None
        
        total_acc = sum([r.get("percentage", 0) for r in rows]) / len(rows)
        
        topic_scores = {}
        for r in rows:
            t = r.get("topic", "Unknown")
            if t not in topic_scores: topic_scores[t] = []
            topic_scores[t].append(r.get("percentage", 0))
            
        weak_topics = []
        for t, scores in topic_scores.items():
            if sum(scores) / len(scores) < 60:
                weak_topics.append(t)
                
        return {
            "accuracy": total_acc,
            "weak_topics": ", ".join(weak_topics) if weak_topics else None
        }
    except Exception as e:
        print("Performance stats error:", e)
        return None

# ----------------------------
# API Routes
# ----------------------------

class UserSignup(BaseModel):
    name: str
    email: str
    password: str

class UserLogin(BaseModel):
    email: str
    password: str

@router.post("/signup")
async def signup(user: UserSignup):
    if users_collection is None:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    existing = users_collection.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")

    hashed_pw = bcrypt.hashpw(user.password.encode('utf-8'), bcrypt.gensalt())
    new_user = {
        "name": user.name,
        "email": user.email,
        "password": hashed_pw.decode('utf-8'),
        "created_at": datetime.datetime.utcnow()
    }
    res = users_collection.insert_one(new_user)
    return {"user_id": str(res.inserted_id), "name": user.name}

@router.post("/login")
async def login(user: UserLogin):
    if users_collection is None:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    record = users_collection.find_one({"email": user.email})
    if not record:
        raise HTTPException(status_code=400, detail="Invalid email or password")
        
    if not bcrypt.checkpw(user.password.encode('utf-8'), record["password"].encode('utf-8')):
        raise HTTPException(status_code=400, detail="Invalid email or password")
        
    return {"user_id": str(record["_id"]), "name": record.get("name", "User")}

@router.post("/generate-quiz")
async def generate_quiz(
    file: UploadFile = File(...),
    topic: str = Form(...),
    difficulty: str = Form(...),
    count: int = Form(...),
    x_user_id: str = Header(None)
):
    if not x_user_id:
        raise HTTPException(status_code=401, detail="Unauthorized: Please log in first.")

    suffix = os.path.splitext(file.filename)[1].lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    if suffix == ".pdf":
        text = extract_pdf(tmp_path)
    elif suffix in [".docx", ".doc"]:
        text = extract_docx(tmp_path)
    elif suffix in [".png", ".jpg", ".jpeg"]:
        text = extract_image(tmp_path)
    else:
        os.remove(tmp_path)
        return {"error": "Unsupported file format"}
    
    os.remove(tmp_path)

    try:
        pipe = process_input(text)

        # Fetch History to power Adaptive mode
        stats = get_user_performance_stats(user_id=x_user_id)

        questions = generate_questions(
            topic,
            difficulty,
            count,
            pipe["content"],
            history_stats=stats
        )

        return {
            "questions": questions,
            "word_count": pipe["word_count"],
            "chunks": pipe["num_chunks"]
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

class SubmitQuizData(BaseModel):
    questions: list
    answers: list
    topic: str

@router.post("/submit-quiz")
async def submit_quiz(data: SubmitQuizData, x_user_id: str = Header(None)):
    if not x_user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    questions = data.questions
    answers = data.answers
    topic = data.topic

    result = analyze_performance(questions, answers, topic)
    
    percentage = (result["score"] / result["total"]) * 100 if result["total"] > 0 else 0

    # Save to history DB
    if history_collection is not None:
        history_collection.insert_one({
            "user_id": x_user_id,
            "topic": topic,
            "score": result["score"],
            "total": result["total"],
            "grade": result["grade"],
            "percentage": percentage,
            "timestamp": datetime.datetime.utcnow()
        })

    return result

@router.get("/history")
async def get_history(x_user_id: str = Header(None)):
    if history_collection is None:
        return {"history": []}
    if not x_user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")
    
    rows = list(history_collection.find({"user_id": x_user_id}).sort("timestamp", -1))
    
    history_list = []
    for row in rows:
        row["id"] = str(row["_id"])
        del row["_id"]
        # Format timestamp to ISO string for frontend if it's a datetime object
        if isinstance(row.get("timestamp"), datetime.datetime):
            row["timestamp"] = row["timestamp"].isoformat() + "Z"
        history_list.append(row)
        
    return {"history": history_list}